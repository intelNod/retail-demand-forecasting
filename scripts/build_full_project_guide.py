"""Build the Russian full-project guide DOCX from its reviewed Markdown source."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image


PROJECT_ROOT = Path(__file__).resolve().parents[1]
SOURCE_PATH = PROJECT_ROOT / "docs" / "FULL_PROJECT_GUIDE_RU.md"
OUTPUT_PATH = PROJECT_ROOT / "submission" / "FULL_PROJECT_GUIDE_RU.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
PALE_BLUE = "F4F7FB"
DARK_TEXT = "263238"
MUTED = "5F6B7A"
WHITE = "FFFFFF"
LIGHT_GRAY = "F3F5F7"
MID_GRAY = "D5DCE3"
GREEN = "70AD47"
ORANGE = "ED7D31"


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def keep_table_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120,
                     bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start),
                          ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = MID_GRAY, size: str = "4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_paragraph_box(paragraph, fill: str, border: str | None = None) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    if border:
        p_bdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), border)
        p_bdr.append(left)
        p_pr.append(p_bdr)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Страница ")
    run.font.name = "Calibri"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])


def add_hyperlink(paragraph, label: str, url: str) -> None:
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_props.extend([color, underline])
    text = OxmlElement("w:t")
    text.text = label
    run.extend([run_props, text])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_PATTERN = re.compile(
    r"(\*\*.+?\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\)|https?://[^\s]+)"
)


def add_inline(paragraph, text: str, *, default_size: float | None = None,
               default_color: str | None = None) -> None:
    position = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position:match.start()])
            if default_size:
                run.font.size = Pt(default_size)
            if default_color:
                run.font.color.rgb = RGBColor.from_string(default_color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
            if default_size:
                run.font.size = Pt(default_size)
            if default_color:
                run.font.color.rgb = RGBColor.from_string(default_color)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(default_size or 9.5)
            run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
        elif token.startswith("["):
            label, url = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token).groups()
            add_hyperlink(paragraph, label, url)
        else:
            clean_url = token.rstrip(".,;)")
            add_hyperlink(paragraph, clean_url, clean_url)
            suffix = token[len(clean_url):]
            if suffix:
                paragraph.add_run(suffix)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        if default_size:
            run.font.size = Pt(default_size)
        if default_color:
            run.font.color.rgb = RGBColor.from_string(default_color)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(DARK_TEXT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    for name, size, color, before, after in (
        ("Title", 30, DARK_BLUE, 0, 14),
        ("Subtitle", 15, MUTED, 0, 12),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = document.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.widow_control = True

    for name in ("List Bullet", "List Number"):
        style = document.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.font.color.rgb = RGBColor.from_string(DARK_TEXT)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    caption = document.styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(10)
    caption.paragraph_format.keep_with_next = False

    header = section.header
    header_p = header.paragraphs[0]
    header_p.text = "RETAIL DEMAND FORECASTING  ·  ПОЛНЫЙ ГАЙД"
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header_p.runs[0]
    header_run.font.name = "Calibri"
    header_run.font.size = Pt(8)
    header_run.font.bold = True
    header_run.font.color.rgb = RGBColor.from_string(BLUE)

    footer = section.footer
    footer_table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    footer_table.columns[0].width = Inches(4.9)
    footer_table.columns[1].width = Inches(1.6)
    left = footer_table.cell(0, 0).paragraphs[0]
    left.text = "Nosirov Nodir · Tulpar · RET-01"
    left.runs[0].font.name = "Calibri"
    left.runs[0].font.size = Pt(8.5)
    left.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
    add_page_number(footer_table.cell(0, 1).paragraphs[0])

    properties = document.core_properties
    properties.title = "Retail Demand Forecasting — полный пошаговый гайд"
    properties.subject = "AI/ML Fundamentals capstone project and defense guide"
    properties.author = "Nosirov Nodir"
    properties.keywords = "RET-01, retail, demand forecasting, M5, machine learning"


def add_cover(document: Document) -> None:
    rule = document.add_table(rows=1, cols=1)
    rule.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = rule.cell(0, 0)
    shade_cell(cell, BLUE)
    set_cell_margins(cell, 0, 0, 0, 0)
    cell.height = Inches(0.08)
    cell.paragraphs[0].paragraph_format.space_after = Pt(0)

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(42)

    label = document.add_paragraph()
    label.paragraph_format.space_after = Pt(8)
    run = label.add_run("AI/ML FUNDAMENTALS · CAPSTONE PROJECT")
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(BLUE)

    title = document.add_paragraph(style="Title")
    title.add_run("Retail Demand\nForecasting")

    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.add_run("Полный пошаговый гайд по проекту и защите")

    badge = document.add_paragraph()
    badge.paragraph_format.space_before = Pt(5)
    badge.paragraph_format.space_after = Pt(34)
    badge_run = badge.add_run("RET-01   ·   FIELD-BASED SCENARIO   ·   RETAILTECH")
    badge_run.font.name = "Calibri"
    badge_run.font.size = Pt(10)
    badge_run.font.bold = True
    badge_run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    set_paragraph_box(badge, LIGHT_BLUE)

    metadata = document.add_table(rows=4, cols=2)
    metadata.alignment = WD_TABLE_ALIGNMENT.LEFT
    metadata.autofit = False
    labels = ("Студент", "Группа", "Репозиторий", "Версия")
    values = (
        "Nosirov Nodir",
        "Tulpar",
        "github.com/intelNod/retail-demand-forecasting",
        "1 августа 2026",
    )
    for row, label_text, value_text in zip(metadata.rows, labels, values):
        row.cells[0].width = Inches(1.25)
        row.cells[1].width = Inches(5.25)
        for cell_item in row.cells:
            set_cell_margins(cell_item, 40, 0, 40, 80)
        label_p = row.cells[0].paragraphs[0]
        label_run = label_p.add_run(label_text.upper())
        label_run.font.name = "Calibri"
        label_run.font.size = Pt(8.5)
        label_run.font.bold = True
        label_run.font.color.rgb = RGBColor.from_string(MUTED)
        value_p = row.cells[1].paragraphs[0]
        if label_text == "Репозиторий":
            add_hyperlink(value_p, value_text, "https://" + value_text)
        else:
            value_run = value_p.add_run(value_text)
            value_run.font.name = "Calibri"
            value_run.font.size = Pt(10.5)
            value_run.font.color.rgb = RGBColor.from_string(DARK_TEXT)

    document.add_paragraph().paragraph_format.space_after = Pt(20)
    metrics = document.add_table(rows=1, cols=3)
    metrics.alignment = WD_TABLE_ALIGNMENT.CENTER
    metrics.autofit = False
    for cell_item, number, label_text, color in zip(
        metrics.rows[0].cells,
        ("4.4179", "7.72%", "3.119 M"),
        ("FINAL TEST MAE", "BETTER THAN BASELINE", "WEEKLY ROWS"),
        (BLUE, GREEN, ORANGE),
    ):
        cell_item.width = Inches(2.16)
        set_cell_margins(cell_item, 100, 100, 90, 100)
        shade_cell(cell_item, PALE_BLUE)
        p = cell_item.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        number_run = p.add_run(number + "\n")
        number_run.font.name = "Calibri"
        number_run.font.size = Pt(18)
        number_run.font.bold = True
        number_run.font.color.rgb = RGBColor.from_string(color)
        label_run = p.add_run(label_text)
        label_run.font.name = "Calibri"
        label_run.font.size = Pt(7.5)
        label_run.font.bold = True
        label_run.font.color.rgb = RGBColor.from_string(MUTED)

    document.add_paragraph().paragraph_format.space_after = Pt(22)
    note = document.add_paragraph()
    note.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note_run = note.add_run(
        "От полного аудита M5 до leakage-safe модели, API, UI и защиты"
    )
    note_run.font.name = "Calibri"
    note_run.font.size = Pt(9.5)
    note_run.font.italic = True
    note_run.font.color.rgb = RGBColor.from_string(MUTED)

    document.add_page_break()


def is_block_start(line: str) -> bool:
    stripped = line.strip()
    return bool(
        not stripped
        or stripped.startswith(("#", ">", "```", "![", "|"))
        or stripped == "---"
        or re.match(r"^\s*[-*] ", line)
        or re.match(r"^\s*\d+\. ", line)
        or (stripped.startswith("*") and stripped.endswith("*") and len(stripped) > 2)
    )


def add_heading(document: Document, level: int, text: str) -> None:
    heading = document.add_heading(level=min(level, 3))
    heading.paragraph_format.keep_with_next = True
    if level == 1:
        heading.paragraph_format.page_break_before = True
    add_inline(heading, text)


def add_quote(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.24)
    paragraph.paragraph_format.right_indent = Inches(0.1)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.keep_together = True
    add_inline(paragraph, text, default_color=DARK_BLUE)
    for run in paragraph.runs:
        run.italic = True
    set_paragraph_box(paragraph, PALE_BLUE, BLUE)


def add_code(document: Document, code_lines: list[str]) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.2)
    paragraph.paragraph_format.right_indent = Inches(0.1)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.keep_together = True
    run = paragraph.add_run("\n".join(code_lines))
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(DARK_TEXT)
    set_paragraph_box(paragraph, LIGHT_GRAY)


def add_horizontal_rule(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(5)
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), MID_GRAY)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def parse_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def add_markdown_table(document: Document, lines: list[str]) -> None:
    rows = [parse_table_row(line) for line in lines]
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", cell) for cell in rows[1]):
        rows.pop(1)
    column_count = max(len(row) for row in rows)
    rows = [row + [""] * (column_count - len(row)) for row in rows]
    table = document.add_table(rows=len(rows), cols=column_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    tbl_pr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)

    max_lengths = []
    for column_index in range(column_count):
        maximum = max(6, max(len(row[column_index]) for row in rows))
        max_lengths.append(min(maximum, 42))
    total = sum(max_lengths)
    widths = [max(0.65, 6.5 * length / total) for length in max_lengths]
    width_total = sum(widths)
    widths = [width * 6.5 / width_total for width in widths]

    for row_index, (table_row, values) in enumerate(zip(table.rows, rows)):
        keep_table_row_together(table_row)
        if row_index == 0:
            set_repeat_table_header(table_row)
        for column_index, (cell, value) in enumerate(zip(table_row.cells, values)):
            cell.width = Inches(widths[column_index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            shade_cell(cell, LIGHT_BLUE if row_index == 0 else (PALE_BLUE if row_index % 2 == 0 else WHITE))
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.1
            add_inline(paragraph, value, default_size=8.5)
            if row_index == 0:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    after = document.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def add_image(document: Document, source_dir: Path, line: str) -> None:
    match = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", line.strip())
    if not match:
        raise ValueError(f"Invalid Markdown image: {line}")
    alt_text, relative_path = match.groups()
    image_path = (source_dir / relative_path).resolve()
    if not image_path.exists():
        raise FileNotFoundError(image_path)
    with Image.open(image_path) as image_file:
        width_px, height_px = image_file.size
    max_width, max_height = 6.25, 7.1
    aspect = width_px / height_px
    width = min(max_width, max_height * aspect)
    height = width / aspect
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run()
    shape = run.add_picture(str(image_path), width=Inches(width), height=Inches(height))
    shape._inline.docPr.set("descr", alt_text)


def add_list(document: Document, lines: list[str], numbered: bool) -> None:
    marker_pattern = r"^\s*\d+\.\s+" if numbered else r"^\s*[-*]\s+"
    items: list[str] = []
    current = ""
    for line in lines:
        if re.match(marker_pattern, line):
            if current:
                items.append(current.strip())
            current = re.sub(marker_pattern, "", line).strip()
        else:
            current += " " + line.strip()
    if current:
        items.append(current.strip())

    for item in items:
        is_checkbox = item.startswith("[ ] ") or item.startswith("[x] ")
        if is_checkbox:
            checked = item.startswith("[x] ")
            item = item[4:]
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.375)
            paragraph.paragraph_format.first_line_indent = Inches(-0.188)
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.add_run("☒ " if checked else "☐ ")
            add_inline(paragraph, item)
        else:
            paragraph = document.add_paragraph(style="List Number" if numbered else "List Bullet")
            add_inline(paragraph, item)


def build_body(document: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    first_rule = lines.index("---")
    lines = lines[first_rule + 1:]
    index = 0
    source_dir = SOURCE_PATH.parent

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if not stripped:
            index += 1
            continue
        if stripped == "---":
            next_index = index + 1
            while next_index < len(lines) and not lines[next_index].strip():
                next_index += 1
            if not (
                next_index < len(lines)
                and re.match(r"^#\s+", lines[next_index].strip())
            ):
                add_horizontal_rule(document)
            index += 1
            continue
        if stripped.startswith("```"):
            code_lines: list[str] = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            index += 1
            add_code(document, code_lines)
            continue
        if stripped.startswith("#"):
            match = re.match(r"^(#{1,3})\s+(.+)$", stripped)
            if not match:
                raise ValueError(f"Invalid heading: {line}")
            add_heading(document, len(match.group(1)), match.group(2))
            index += 1
            continue
        if stripped.startswith(">"):
            quote_lines: list[str] = []
            while index < len(lines) and lines[index].strip().startswith(">"):
                quote_lines.append(lines[index].strip()[1:].strip())
                index += 1
            add_quote(document, " ".join(quote_lines))
            continue
        if stripped.startswith("!["):
            add_image(document, source_dir, stripped)
            index += 1
            continue
        if stripped.startswith("|"):
            table_lines: list[str] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_markdown_table(document, table_lines)
            continue
        if re.match(r"^\s*[-*]\s+", line):
            list_lines: list[str] = []
            while index < len(lines):
                candidate = lines[index]
                if re.match(r"^\s*[-*]\s+", candidate) or (
                    list_lines and candidate.startswith(("  ", "\t")) and candidate.strip()
                ):
                    list_lines.append(candidate)
                    index += 1
                else:
                    break
            add_list(document, list_lines, numbered=False)
            continue
        if re.match(r"^\s*\d+\.\s+", line):
            list_lines = []
            while index < len(lines):
                candidate = lines[index]
                if re.match(r"^\s*\d+\.\s+", candidate) or (
                    list_lines and candidate.startswith(("  ", "\t")) and candidate.strip()
                ):
                    list_lines.append(candidate)
                    index += 1
                else:
                    break
            add_list(document, list_lines, numbered=True)
            continue
        if stripped.startswith("*") and stripped.endswith("*") and len(stripped) > 2:
            paragraph = document.add_paragraph(style="Caption")
            add_inline(paragraph, stripped[1:-1])
            index += 1
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines) and not is_block_start(lines[index]):
            paragraph_lines.append(lines[index].strip())
            index += 1
        paragraph = document.add_paragraph()
        add_inline(paragraph, " ".join(paragraph_lines))


def main() -> None:
    markdown = SOURCE_PATH.read_text(encoding="utf-8")
    if markdown.count("![") != 10:
        raise RuntimeError("The reviewed guide must contain exactly 10 figures.")
    document = Document()
    configure_document(document)
    add_cover(document)
    build_body(document, markdown)
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
